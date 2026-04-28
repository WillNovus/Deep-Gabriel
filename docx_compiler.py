import io
import os
import re
import ast
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Tuple, Optional, Callable
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from src.my_agent.utils.docx_math import DocxMathBuilder
except ImportError:
    class DummyMathBuilder:
        def latex_to_omml(self, clean_latex): return None
    DocxMathBuilder = DummyMathBuilder

logger = logging.getLogger(__name__)

# Resolve logo once at module load — same directory as this file
_LOGO_PATH = Path(__file__).resolve().parent / "Deep-Gabriel.png"

def _add_formatted_run(paragraph, text):
    """
    Parses **bold** and *italic* within a text string and adds runs to the paragraph.
    """
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if not part: 
            continue
        
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def add_smart_paragraph(doc, text, math_builder, style=None):
    """
    Adds a paragraph that handles:
    1. LaTeX Math ($...$) -> Native Word Equation (OMML)
    2. Markdown Formatting (**bold**, *italic*) -> Rich Text
    """
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()

    parts = re.split(r'(\$[^$]+\$)', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('$') and part.endswith('$') and len(part) > 2 and math_builder:
            clean_latex = part[1:-1].strip()
            omml_element = math_builder.latex_to_omml(clean_latex)
            
            if omml_element is not None:
                p.element.append(omml_element)
            else:
                p.add_run(part)
        else:
            _add_formatted_run(p, part)

def generate_docx_from_state(
    state: dict,
    image_fetcher: Optional[Callable[[str], Optional[bytes]]] = None
) -> Tuple[str, bytes]:
    """
    📄 Generates a formatted .docx file with native math support.
    Order: Title -> Body -> References -> Appendix (Diagrams)

    Args:
        state: Dictionary containing final_paper, bibliography, diagram_manifest, etc.
        image_fetcher: Optional callback to fetch image bytes by filename from MinIO.
                      Signature: (filename: str) -> Optional[bytes]

    Returns:
        Tuple of (filename, docx_bytes) for upload to MinIO.
    """
    math_builder = DocxMathBuilder()

    raw_topic = state.get("topic", "Untitled Research Topic")
    if isinstance(raw_topic, str) and ("type" in raw_topic or "text" in raw_topic):
        try:
            parsed = ast.literal_eval(raw_topic)
            if isinstance(parsed, list) and len(parsed) > 0:
                raw_topic = parsed[0].get("text", "Untitled")
        except Exception:
            match = re.search(r"'text':\s*['\"](.*?)['\"]", raw_topic)
            if match:
                raw_topic = match.group(1)
    topic = str(raw_topic).strip().strip('"').strip("'")
    citation_style = state.get("citation_style", "APA")
    final_paper = state.get("final_paper", "") or state.get("final_draft", "")
    bibliography = state.get("bibliography", {})
    manifest = state.get("diagram_manifest", {})

    doc = Document()

    # --- PAGE SETUP ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # --- HEADER LOGO (upper right) ---
        if _LOGO_PATH.exists():
            try:
                header = section.header
                header.is_linked_to_previous = False
                para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = para.add_run()
                run.add_picture(str(_LOGO_PATH), width=Inches(0.5))
            except Exception as e:
                logger.warning("Could not embed header logo: %s", e)
        else:
            logger.warning("Logo not found at %s — skipping header logo", _LOGO_PATH)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    try:
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except AttributeError:
        pass
    style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)

    doc.add_heading(topic, level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Citation Style: {citation_style}")
    doc.add_paragraph("")

    if not final_paper or not final_paper.strip():
        doc.add_paragraph("No content provided.")
    else:
        lines = final_paper.split('\n')

        for line in lines:
            line = line.strip()
            if not line: continue

            if line.startswith("#"):
                level = len(line.split(' ')[0])
                clean_text = line.lstrip('#').strip()
                clean_text = clean_text.replace('**', '')
                doc.add_heading(clean_text, level=min(level, 3))

            elif line.startswith('**') and line.endswith('**') and len(line) < 100:
                clean_text = line[2:-2]
                doc.add_heading(clean_text, level=2)

            elif line.startswith("- ") or line.startswith("* "):
                add_smart_paragraph(doc, line[2:], math_builder, style='List Bullet')

            else:
                add_smart_paragraph(doc, line, math_builder)

    if bibliography:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        refs = bibliography.values() if isinstance(bibliography, dict) else bibliography
        for entry in refs:
            text = entry.get("formatted", "") if isinstance(entry, dict) else str(entry)
            if text:
                doc.add_paragraph(text)

    # Handle diagrams - fetch from MinIO if image_fetcher provided
    if manifest:
        doc.add_page_break()
        doc.add_heading("Appendix: Visualizations", level=1)

        for section, data in manifest.items():
            doc.add_heading(f"Figure: {data['caption']}", level=3)

            # Get filename - prefer explicit 'filename' field, fallback to extracting from 'path'
            image_filename = data.get('filename') or os.path.basename(data.get('path', ''))

            if image_filename:
                image_bytes = None

                # Try to fetch image from MinIO via callback
                if image_fetcher:
                    try:
                        image_bytes = image_fetcher(image_filename)
                    except Exception as e:
                        logger.warning("Failed to fetch image %s from MinIO: %s", image_filename, e)

                if image_bytes:
                    # Use temporary file for python-docx (requires file path or file-like object)
                    try:
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            tmp.write(image_bytes)
                            tmp_path = tmp.name

                        doc.add_picture(tmp_path, width=Inches(6))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1  # Center
                        logger.info("Embedded diagram for %s from MinIO", section)

                        # Clean up temp file
                        os.unlink(tmp_path)
                    except Exception as e:
                        logger.error("Docx insertion error: %s", e)
                        doc.add_paragraph(f"[Error rendering image: {e}]")
                else:
                    logger.warning("Image not found in MinIO: %s", image_filename)
                    doc.add_paragraph(f"[Image not available: {image_filename}]")

    # Generate filename
    short_topic = "_".join(topic.split()[:6])
    safe_topic = re.sub(r'\W+', '_', short_topic.lower())
    filename = f"research_{safe_topic}.docx"

    # Save to BytesIO and return bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()

    logger.info("Generated DOCX: %s (%d bytes)", filename, len(docx_bytes))
    return filename, docx_bytes


def compose_markdown_preview(draft: str, diagram_manifest: dict) -> str:
    """
    Composes a markdown string with inline Mermaid code blocks embedded
    at their target section positions.

    Args:
        draft: The markdown draft text.
        diagram_manifest: Dict mapping section titles to diagram metadata
                          (must include 'mermaid_code' and 'caption' keys).

    Returns:
        Markdown string with ```mermaid blocks inserted after matching headings.
    """
    if not diagram_manifest:
        return draft

    # Build a list of diagrams that have mermaid source code
    diagrams = []
    for section_title, data in diagram_manifest.items():
        mermaid_code = data.get("mermaid_code")
        if mermaid_code:
            diagrams.append({
                "section": section_title,
                "code": mermaid_code,
                "caption": data.get("caption", section_title),
            })

    if not diagrams:
        return draft

    lines = draft.split("\n")
    result_lines = []
    placed_sections = set()

    for i, line in enumerate(lines):
        result_lines.append(line)

        # Check if this line is a heading that matches a diagram target section
        stripped = line.strip()
        if not stripped.startswith("#"):
            continue

        heading_text = stripped.lstrip("#").strip()

        for diag in diagrams:
            if diag["section"] in placed_sections:
                continue
            # Case-insensitive substring match — the LLM-generated section title
            # may not exactly match the heading text
            if (diag["section"].lower() in heading_text.lower()
                    or heading_text.lower() in diag["section"].lower()):
                # Insert mermaid block after the heading
                result_lines.append("")
                result_lines.append(f"```mermaid")
                result_lines.append(diag["code"])
                result_lines.append(f"```")
                result_lines.append(f"*Figure: {diag['caption']}*")
                result_lines.append("")
                placed_sections.add(diag["section"])
                break

    # Append any unplaced diagrams at the end
    unplaced = [d for d in diagrams if d["section"] not in placed_sections]
    if unplaced:
        result_lines.append("")
        result_lines.append("## Appendix: Diagrams")
        result_lines.append("")
        for diag in unplaced:
            result_lines.append(f"### {diag['caption']}")
            result_lines.append("")
            result_lines.append(f"```mermaid")
            result_lines.append(diag["code"])
            result_lines.append(f"```")
            result_lines.append(f"*Figure: {diag['caption']}*")
            result_lines.append("")

    return "\n".join(result_lines)